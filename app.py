"""
高中物理学习规划生成器 - Web版
基于美博教育20+年经验，集成AI模型自动生成Excel课规表和Word辅导方案
"""
import os
import io
import json
import base64
import time
import uuid
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
import requests
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
CORS(app)

# ============ 配置 ============
DEFAULT_AI_URL = "http://110.185.163.23:50000/v1"
DEFAULT_MODEL = "qwen3.6-35b"
API_KEY = os.environ.get("AI_API_KEY", "")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "outputs")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# 美博品牌色
BRAND_GREEN = "2E7D32"
BRAND_LIGHT_GREEN = "E8F5E9"

# ============ 系统提示词 ============
SYSTEM_PROMPT = """你是一位专业的高中物理教育规划专家，基于美博教育二十多年的教学经验，为学生提供科学、系统的高中物理学习规划。

## 核心能力
1. **五维度学情剖析**：基础掌握、学习习惯、应试能力、心理状态、性格特点
2. **课程规划**：按A（基础夯实）→B（深化突破）→C（综合复习）递进设计
3. **题型分值目标**：基于高考物理题型分值结构设定目标

## 高考物理知识体系
- 必修1：运动学（匀变速、图像）、力学（三种力、受力分析、共点力平衡）、牛顿定律（瞬时性、超重失重、连接体、传送带、板块）
- 必修2：曲线运动（平抛、圆周）、万有引力（卫星变轨）、机械能（动能定理、能量守恒）
- 选修3-1：电场、恒定电流、磁场
- 选修3-2：电磁感应、交变电流
- 选修3-3/3-4/3-5：热学、机械振动机械波、光学、原子物理

## 课时分配法则
- 电磁学核心（磁场+电磁感应）：40%+
- 力学核心（牛顿定律+曲线运动+能量动量）：30%
- 热光原（热学+光学+原子物理）：20%
- 实验专题：10%

## 输出要求
你必须以JSON格式返回以下完整结构，不要省略任何字段：

{
  "student_analysis": {
    "basic": "物理基础掌握情况分析，2-3句话",
    "habit": "学习习惯分析，2-3句话",
    "exam": "应试能力分析，2-3句话",
    "psychology": "心理状态分析，1-2句话",
    "personality": "性格特点，1句话",
    "overall_suggestion": "综合建议，2-3句话"
  },
  "teacher_feedback_summary": [
    {"topic": "反馈涉及的知识点", "content": "教师原始反馈摘要", "insight": "教学启示，如何与当前计划衔接"}
  ],
  "course_plan": [
    {"lesson": 1, "module": "模块名", "content": "具体教学内容", "hours": 2, "goal": "本课次目标"}
  ],
  "teaching_approach": {
    "phase1": {"name": "第一阶段名称", "period": "时间段", "focus": "重点内容", "strategy": "教学策略"},
    "phase2": {"name": "第二阶段名称", "period": "时间段", "focus": "重点内容", "strategy": "教学策略"},
    "phase3": {"name": "第三阶段名称", "period": "时间段", "focus": "重点内容", "strategy": "教学策略"}
  },
  "expected_goals": {
    "single_choice": {"score": 28, "target": 20},
    "multi_choice": {"score": 18, "target": 12},
    "experiment": {"score": 16, "target": 12},
    "question_13": {"score": 10, "target": 8},
    "question_14": {"score": 12, "target": 8},
    "question_15": {"score": 16, "target": 10},
    "total_target": 70
  },
  "suggestions": ["学习建议1", "学习建议2", "学习建议3"]
}

注意：
- course_plan的lesson从1开始顺序编号，每课次2小时
- content要具体，写清楚教哪些知识点和题型
- 课程设计遵循A→B→C递进逻辑
- 使用"孩子"称呼学生
- 根据学生的薄弱点调整课程重点
- expected_goals要基于学生当前水平和目标动态调整
- total_target = 各题型target之和"""


def call_ai_model(api_key, ai_url, model, messages):
    """调用AI模型API"""
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    payload = {
        "model": model,
        "messages": messages,
        "temperature": 0.7,
        "max_tokens": 8192
    }
    resp = requests.post(f"{ai_url}/chat/completions", headers=headers, json=payload, timeout=120)
    resp.raise_for_status()
    result = resp.json()
    return result["choices"][0]["message"]["content"]


def extract_json_from_response(text):
    """从AI回复中提取JSON"""
    # 尝试直接解析
    try:
        return json.loads(text)
    except:
        pass

    # 尝试提取```json ... ```块
    import re
    match = re.search(r'```(?:json)?\s*([\s\S]*?)```', text)
    if match:
        try:
            return json.loads(match.group(1))
        except:
            pass

    # 尝试找到第一个{和最后一个}
    start = text.find('{')
    end = text.rfind('}')
    if start != -1 and end != -1:
        try:
            return json.loads(text[start:end+1])
        except:
            pass

    raise ValueError("无法从AI回复中提取有效的JSON数据")


# ============ Excel生成 ============
def generate_excel(student_name, plan_data, output_path):
    """生成Excel课规表"""
    wb = Workbook()
    ws = wb.active
    ws.title = "物理课规表"

    # 样式定义
    title_font = Font(name="微软雅黑", size=16, bold=True, color=BRAND_GREEN)
    header_font = Font(name="微软雅黑", size=11, bold=True, color="FFFFFF")
    cell_font = Font(name="微软雅黑", size=10)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    header_fill = PatternFill(start_color=BRAND_GREEN, end_color=BRAND_GREEN, fill_type="solid")
    alt_fill = PatternFill(start_color=BRAND_LIGHT_GREEN, end_color=BRAND_LIGHT_GREEN, fill_type="solid")

    # 标题行
    ws.merge_cells("A1:E1")
    title_cell = ws["A1"]
    title_cell.value = f"{student_name} 物理课规表"
    title_cell.font = title_font
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 40

    # 表头
    headers = ["课次", "课程模块", "课程内容", "课时安排", "课程目标"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border

    # 数据行
    course_plan = plan_data.get("course_plan", [])
    for i, item in enumerate(course_plan):
        row = i + 3
        values = [
            item.get("lesson", i+1),
            item.get("module", ""),
            item.get("content", ""),
            f"{item.get('hours', 2)}h",
            item.get("goal", "")
        ]
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row, column=col, value=val)
            cell.font = cell_font
            cell.border = thin_border
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            if col <= 2:
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            if i % 2 == 0:
                cell.fill = alt_fill

    # 列宽
    col_widths = [8, 16, 45, 10, 30]
    for col, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(col)].width = w

    # 行高
    for i in range(3, 3 + len(course_plan)):
        ws.row_dimensions[i].height = 45

    wb.save(output_path)
    return output_path


# ============ Word生成 ============
def generate_word(student_name, student_info, plan_data, output_path):
    """生成Word个性化辅导方案"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页脚
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run("美博深度辅导  好学习有动力")
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 封面标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(6)
    tr = title.add_run("美博教育个性化辅导方案")
    tr.font.size = Pt(22)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(f"{student_name} 同学 物理 辅导")
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    school_line = doc.add_paragraph()
    school_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    slr = school_line.add_run(f"{student_info.get('school', '')} | {student_info.get('grade', '')}")
    slr.font.size = Pt(11)
    slr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    city_line = doc.add_paragraph()
    city_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clr = city_line.add_run("成都美博教育")
    clr.font.size = Pt(12)

    doc.add_page_break()

    # 辅助函数
    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
        return h

    def add_table_with_style(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(10)
            # 背景色
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), BRAND_GREEN)
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Data rows
        for r, row_data in enumerate(rows):
            for c, val in enumerate(row_data):
                cell = table.rows[r+1].cells[c]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)
        return table

    # ===== 一、学生情况剖析 =====
    add_heading("一、学生情况剖析", level=1)
    analysis = plan_data.get("student_analysis", {})

    sections_zh = {
        "basic": ("1.1 物理基础掌握情况", "基础掌握"),
        "habit": ("1.2 物理学习习惯", "学习习惯"),
        "exam": ("1.3 物理应试能力", "应试能力"),
        "psychology": ("1.4 物理学习心理状态", "心理状态"),
        "personality": ("1.5 学生性格特点", "性格特点"),
    }

    for key, (title, label) in sections_zh.items():
        content = analysis.get(key, "")
        if content:
            add_heading(title, level=2)
            p = doc.add_paragraph(content)
            p.paragraph_format.first_line_indent = Pt(22)

    # 综合建议
    overall = analysis.get("overall_suggestion", "")
    if overall:
        add_heading("1.6 综合建议", level=2)
        p = doc.add_paragraph(overall)
        p.paragraph_format.first_line_indent = Pt(22)

    # ===== 二、教学基本内容及课时安排 =====
    add_heading("二、教学基本内容及课时安排", level=1)
    course_plan = plan_data.get("course_plan", [])
    course_rows = [[
        str(item.get("lesson", i+1)),
        item.get("module", ""),
        item.get("content", ""),
        f"{item.get('hours', 2)}h",
        item.get("goal", "")
    ] for i, item in enumerate(course_plan)]
    add_table_with_style(["课次", "课程模块", "课程内容", "课时", "课程目标"], course_rows)

    # 总结
    total_hours = sum(item.get("hours", 2) for item in course_plan)
    p = doc.add_paragraph(f"\n共计{len(course_plan)}课次，{total_hours}课时。课程设计遵循基础夯实→深化突破→综合复习的递进逻辑。")

    # ===== 三、个性化教学思路 =====
    add_heading("三、个性化教学思路", level=1)
    phases = plan_data.get("teaching_approach", {})
    for i, (key, phase) in enumerate(phases.items(), 1):
        name = phase.get("name", f"第{i}阶段")
        period = phase.get("period", "")
        focus = phase.get("focus", "")
        strategy = phase.get("strategy", "")
        add_heading(f"3.{i} {name}（{period}）", level=2)
        doc.add_paragraph(f"重点内容：{focus}")
        doc.add_paragraph(f"教学策略：{strategy}")

    # ===== 四、预期目标 =====
    add_heading("四、预期目标", level=1)
    goals = plan_data.get("expected_goals", {})

    goal_headers = ["题型", "满分值", "目标分值"]
    goal_rows = [
        ["单选题（7道）", f"{goals.get('single_choice', {}).get('score', 28)}分", f"{goals.get('single_choice', {}).get('target', 20)}分"],
        ["多选题（3道）", f"{goals.get('multi_choice', {}).get('score', 18)}分", f"{goals.get('multi_choice', {}).get('target', 12)}分"],
        ["实验题（2道）", f"{goals.get('experiment', {}).get('score', 16)}分", f"{goals.get('experiment', {}).get('target', 12)}分"],
        ["计算题13题", f"{goals.get('question_13', {}).get('score', 10)}分", f"{goals.get('question_13', {}).get('target', 8)}分"],
        ["计算题14题", f"{goals.get('question_14', {}).get('score', 12)}分", f"{goals.get('question_14', {}).get('target', 8)}分"],
        ["计算题15题", f"{goals.get('question_15', {}).get('score', 16)}分", f"{goals.get('question_15', {}).get('target', 10)}分"],
        ["总计", "100分", f"{goals.get('total_target', 70)}分"],
    ]
    add_table_with_style(goal_headers, goal_rows)

    # ===== 五、学习建议 =====
    add_heading("五、学习建议与注意事项", level=1)
    suggestions = plan_data.get("suggestions", [])
    for i, s in enumerate(suggestions, 1):
        doc.add_paragraph(f"{i}. {s}", style='List Number')

    # ===== 教师反馈摘录 =====
    feedback_list = plan_data.get("teacher_feedback_summary", [])
    if feedback_list:
        add_heading("附录：近期教师课堂反馈摘录", level=1)
        for fb in feedback_list:
            topic = fb.get("topic", "")
            content = fb.get("content", "")
            insight = fb.get("insight", "")
            if topic:
                p = doc.add_paragraph()
                r = p.add_run(f"👨🏻‍🏫 课堂内容：{topic}")
                r.font.bold = True
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(22)
            p.add_run(f"📚 {content}")
            if insight:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Pt(22)
                r = p.add_run(f"💡 教学启示：{insight}")
                r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    # 保存
    doc.save(output_path)
    return output_path


# ============ API路由 ============

@app.route("/")
def index():
    """主页面"""
    return render_template("index.html")


@app.route("/api/generate", methods=["POST"])
def generate():
    """生成规划文件"""
    try:
        # 获取表单数据
        api_key = request.form.get("api_key", "").strip() or API_KEY
        ai_url = request.form.get("ai_url", "").strip() or DEFAULT_AI_URL

        if not api_key:
            return jsonify({"success": False, "error": "请填写API Key"}), 400

        # 学生信息
        student_name = request.form.get("student_name", "").strip()
        school = request.form.get("school", "").strip()
        grade = request.form.get("grade", "").strip()
        current_progress = request.form.get("current_progress", "").strip()
        plan_time = request.form.get("plan_time", "").strip()
        lesson_count = request.form.get("lesson_count", "").strip()
        total_hours = request.form.get("total_hours", "").strip()
        content_scope = request.form.get("content_scope", "").strip()
        strengths = request.form.get("strengths", "").strip()
        weaknesses = request.form.get("weaknesses", "").strip()
        feedback_text = request.form.get("feedback_text", "").strip()

        # 处理图片
        image_base64_list = []
        image_files = request.files.getlist("feedback_images")
        for img in image_files:
            if img.filename:
                img_data = img.read()
                img_base64 = base64.b64encode(img_data).decode("utf-8")
                image_base64_list.append(img_base64)

        # 构建用户消息
        user_message_parts = [f"""请为以下学生生成高中物理学习规划：

## 学生基本信息
- 姓名：{student_name}
- 学校：{school}
- 年级：{grade}
- 目前进度：{current_progress}
- 规划上课时间：{plan_time}
- 上课节数：{lesson_count}节
- 课时数：{total_hours}课时
- 学习内容范围：{content_scope}
- 学生优点：{strengths}
- 学生缺点：{weaknesses}
"""]

        if feedback_text:
            user_message_parts.append(f"""
## 近期教师课堂反馈（文字）
{feedback_text}
""")

        if image_base64_list:
            user_message_parts.append(f"\n## 近期教师课堂反馈（图片）\n共{len(image_base64_list)}张反馈截图，请仔细阅读图片中的内容。")

        user_message = "\n".join(user_message_parts)

        # 构建消息
        messages = [{"role": "system", "content": SYSTEM_PROMPT}]

        if image_base64_list:
            # 多模态消息
            content_array = [{"type": "text", "text": user_message}]
            for img_b64 in image_base64_list:
                content_array.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:image/png;base64,{img_b64}"}
                })
            messages.append({"role": "user", "content": content_array})
        else:
            messages.append({"role": "user", "content": user_message})

        # 调用AI
        ai_response = call_ai_model(api_key, ai_url, model, messages)
        plan_data = extract_json_from_response(ai_response)

        # 生成唯一ID
        task_id = uuid.uuid4().hex[:8]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # 生成Excel
        excel_path = os.path.join(OUTPUT_DIR, f"{student_name}_{timestamp}_课规表.xlsx")
        generate_excel(student_name, plan_data, excel_path)

        # 生成Word
        student_info = {
            "school": school,
            "grade": grade,
        }
        word_path = os.path.join(OUTPUT_DIR, f"{student_name}_{timestamp}_辅导方案.docx")
        generate_word(student_name, student_info, plan_data, word_path)

        return jsonify({
            "success": True,
            "task_id": task_id,
            "excel_file": os.path.basename(excel_path),
            "word_file": os.path.basename(word_path),
            "excel_url": f"/download/{os.path.basename(excel_path)}",
            "word_url": f"/download/{os.path.basename(word_path)}",
            "plan_data": {
                "student_analysis": plan_data.get("student_analysis", {}),
                "course_plan_summary": len(plan_data.get("course_plan", [])),
                "total_hours": sum(item.get("hours", 2) for item in plan_data.get("course_plan", [])),
                "phases": [p.get("name", "") for p in plan_data.get("teaching_approach", {}).values()],
                "target_score": plan_data.get("expected_goals", {}).get("total_target", 0),
            }
        })

    except requests.exceptions.RequestException as e:
        return jsonify({"success": False, "error": f"AI模型请求失败：{str(e)}。请检查API URL和Key是否正确。"}), 500
    except ValueError as e:
        return jsonify({"success": False, "error": f"AI回复解析失败：{str(e)}。请重试或检查模型是否支持。"}), 500
    except Exception as e:
        return jsonify({"success": False, "error": f"生成失败：{str(e)}"}), 500


@app.route("/download/<filename>")
def download(filename):
    """下载生成的文件"""
    filepath = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "文件不存在或已过期"}), 404

    if filename.endswith(".xlsx"):
        mimetype = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif filename.endswith(".docx"):
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        mimetype = "application/octet-stream"

    return send_file(filepath, mimetype=mimetype, as_attachment=True, download_name=filename)


@app.route("/api/config")
def get_config():
    """获取默认配置"""
    return jsonify({
        "ai_url": DEFAULT_AI_URL,
        "model": DEFAULT_MODEL,
        "has_default_key": bool(API_KEY)
    })


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
